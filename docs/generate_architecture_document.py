from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "architecture_assets"
DOCX_PATH = OUT_DIR / "pulse_architecture_documentation.docx"
PDF_PATH = OUT_DIR / "pulse_architecture_documentation.pdf"

NAVY = RGBColor(11, 42, 91)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 43, 55)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8F1FA"
LIGHT_GREEN = "EAF6EF"
LIGHT_AMBER = "FFF5DB"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
LINE = (75, 91, 110)


def font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill, outline=(90, 104, 118), radius=14, title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)
    f = font(23 if title else 19, bold=title)
    max_width = x2 - x1 - 28
    lines = []
    for part in text.split("\n"):
        words = part.split()
        line = ""
        for word in words:
            trial = f"{line} {word}".strip()
            if draw.textbbox((0, 0), trial, font=f)[2] <= max_width:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = word
        if line:
            lines.append(line)
    line_h = f.size + 6
    total_h = len(lines) * line_h
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=f)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, fill=(24, 34, 45), font=f)
        y += line_h


def arrow(draw, start, end, color=LINE, width=3):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 14, ey - 8), (ex - direction * 14, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - direction * 14), (ex + 8, ey - direction * 14)]
    draw.polygon(points, fill=color)


def diagram(title, boxes, arrows_, filename):
    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 84), fill=(11, 42, 91))
    draw.text((50, 24), title, fill="white", font=font(32, bold=True))
    for start, end in arrows_:
        arrow(draw, start, end)
    for box in boxes:
        draw_box(draw, *box)
    path = ASSET_DIR / filename
    img.save(path, quality=95)
    return path


def build_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {}

    diagrams["high_level"] = diagram(
        "High-Level System Architecture",
        [
            ((70, 150, 320, 250), "Biomedical Technician", "#EAF6EF", None, 14, True),
            ((440, 135, 735, 265), "Flutter App\nScreens + Service Layer", "#E8F1FA", None, 14, True),
            ((870, 110, 1180, 215), "Supabase\nAuth, DB, Storage", "#FFF5DB", None, 14, True),
            ((870, 275, 1180, 380), "Local SQLite\nOffline Cache", "#F3F5F7", None, 14, True),
            ((1270, 135, 1530, 265), "Google Sign-In", "#F3F5F7", None, 14, True),
            ((440, 500, 735, 630), "Flask Backend API\nProtected AI/RAG", "#E8F1FA", None, 14, True),
            ((870, 500, 1180, 630), "RAG Service\nManual Retrieval", "#EAF6EF", None, 14, True),
            ((1270, 500, 1530, 630), "Gemini API\nAI Responses", "#FFF5DB", None, 14, True),
        ],
        [
            ((320, 200), (440, 200)),
            ((735, 180), (870, 160)),
            ((735, 230), (870, 325)),
            ((1180, 160), (1270, 190)),
            ((590, 265), (590, 500)),
            ((735, 565), (870, 565)),
            ((1180, 565), (1270, 565)),
        ],
        "01_high_level_architecture.png",
    )

    diagrams["auth"] = diagram(
        "Authentication Flow",
        [
            ((70, 170, 310, 270), "User taps\nGoogle Sign-In", "#EAF6EF", None, 14, True),
            ((410, 170, 650, 270), "Native Google\nSign-In Popup", "#F3F5F7", None, 14, True),
            ((750, 170, 990, 270), "ID Token +\nAccess Token", "#FFF5DB", None, 14, True),
            ((1090, 170, 1360, 270), "Supabase\nsignInWithIdToken", "#E8F1FA", None, 14, True),
            ((1090, 410, 1360, 510), "Supabase Session\nJWT Access Token", "#EAF6EF", None, 14, True),
            ((750, 410, 990, 510), "Upsert User\nProfile", "#F3F5F7", None, 14, True),
            ((410, 410, 650, 510), "Authenticated\nHome Screen", "#E8F1FA", None, 14, True),
        ],
        [
            ((310, 220), (410, 220)),
            ((650, 220), (750, 220)),
            ((990, 220), (1090, 220)),
            ((1225, 270), (1225, 410)),
            ((1090, 460), (990, 460)),
            ((750, 460), (650, 460)),
        ],
        "02_authentication_flow.png",
    )

    diagrams["data"] = diagram(
        "Data and Offline-First Architecture",
        [
            ((75, 170, 360, 280), "Flutter Screens\nAssets, Logs, Manuals", "#E8F1FA", None, 14, True),
            ((500, 170, 800, 280), "DatabaseHelper\nSingle Data Gateway", "#EAF6EF", None, 14, True),
            ((950, 105, 1280, 210), "Supabase Tables\nCloud Records", "#FFF5DB", None, 14, True),
            ((950, 260, 1280, 365), "Supabase Storage\nImages + PDFs", "#FFF5DB", None, 14, True),
            ((950, 450, 1280, 555), "SQLite Cache\nOffline Reads/Writes", "#F3F5F7", None, 14, True),
            ((500, 450, 800, 555), "Sync/Fallback Logic\nNetwork-Aware", "#E8F1FA", None, 14, True),
        ],
        [
            ((360, 225), (500, 225)),
            ((800, 205), (950, 160)),
            ((800, 245), (950, 310)),
            ((650, 280), (650, 450)),
            ((800, 502), (950, 502)),
        ],
        "03_data_architecture.png",
    )

    diagrams["rag"] = diagram(
        "AI and RAG Request Flow",
        [
            ((70, 160, 300, 260), "Technician\nasks question", "#EAF6EF", None, 14, True),
            ((390, 160, 650, 260), "Flutter App\nBearer JWT", "#E8F1FA", None, 14, True),
            ((740, 160, 1000, 260), "Flask API\nverify token", "#F3F5F7", None, 14, True),
            ((1090, 160, 1370, 260), "RAG Service\nretrieve context", "#EAF6EF", None, 14, True),
            ((1090, 420, 1370, 520), "Manual Chunks\nSources", "#FFF5DB", None, 14, True),
            ((740, 420, 1000, 520), "Gemini API\nanswer generation", "#FFF5DB", None, 14, True),
            ((390, 420, 650, 520), "Answer + Sources\nshown in app", "#E8F1FA", None, 14, True),
        ],
        [
            ((300, 210), (390, 210)),
            ((650, 210), (740, 210)),
            ((1000, 210), (1090, 210)),
            ((1230, 260), (1230, 420)),
            ((1090, 470), (1000, 470)),
            ((740, 470), (650, 470)),
        ],
        "04_ai_rag_flow.png",
    )

    diagrams["deployment"] = diagram(
        "Android Google Sign-In Release Build Chain",
        [
            ((70, 155, 330, 255), "Flutter Release\nBuild APK", "#E8F1FA", None, 14, True),
            ((430, 155, 690, 255), "Release Keystore\nSigns APK", "#EAF6EF", None, 14, True),
            ((790, 155, 1050, 255), "SHA-1 / SHA-256\nFingerprints", "#FFF5DB", None, 14, True),
            ((1150, 155, 1450, 255), "Firebase Android\nOAuth Client", "#F3F5F7", None, 14, True),
            ((790, 420, 1050, 520), "google-services.json\nDownloaded", "#F3F5F7", None, 14, True),
            ((430, 420, 690, 520), "Supabase Google\nProvider Config", "#E8F1FA", None, 14, True),
            ((70, 420, 330, 520), "Installed APK\nGoogle Auth Works", "#EAF6EF", None, 14, True),
        ],
        [
            ((330, 205), (430, 205)),
            ((690, 205), (790, 205)),
            ((1050, 205), (1150, 205)),
            ((1300, 255), (930, 420)),
            ((790, 470), (690, 470)),
            ((430, 470), (330, 470)),
        ],
        "05_deployment_auth_chain.png",
    )
    return diagrams


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.bold = bold
    run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Title", 25, NAVY, 0, 8),
        ("Subtitle", 12, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Pulse Architecture Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def para(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_diagram(doc, title, path):
    para(doc, title, "Heading 3")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.55))


def component_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1.55, 2.35, 2.3]
    headers = ["Layer", "Main Files", "Responsibility"]
    for idx, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, headers[idx], bold=True, color=NAVY)
        cell.width = Inches(widths[idx])

    rows = [
        ("App shell", "main.dart", "Initialize Supabase, validate session, route to login or home."),
        ("Auth", "auth_service.dart", "Email auth, Google auth, password reset, session cleanup, user profile sync."),
        ("Data", "database_helper.dart", "SQLite cache, Supabase CRUD, storage image/PDF handling, offline fallback."),
        ("AI/RAG", "rag_api_service.dart, backend/app.py, backend/rag_service.py", "Protected manual retrieval, source-backed answers, Gemini integration."),
        ("Collaboration", "chat_service.dart, google_chat_service.dart", "Team communication, meeting links, Google Chat notifications."),
        ("Operations", "notification_service.dart, sync_service.dart", "Local notifications, background behavior, syncing strategy."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(idx == 0), color=NAVY if idx == 0 else DARK)
            cells[idx].width = Inches(widths[idx])
    return table


def build_docx():
    diagrams = build_diagrams()
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Pulse Clinical Equipment Maintenance Platform")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Architecture and Crucial Implementation Documentation")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared as a technical project documentation brief")
    r.font.color.rgb = MUTED
    r.font.size = Pt(10.5)

    para(
        doc,
        "Pulse is a Flutter-based clinical equipment maintenance platform for biomedical technicians. "
        "It combines mobile-first workflows, authenticated cloud services, offline storage, AI diagnostics, "
        "manual retrieval, and collaboration tools into one operational system.",
    )

    add_diagram(doc, "Figure 1. System architecture at a glance", diagrams["high_level"])

    para(doc, "1. Architecture Overview", "Heading 1")
    para(
        doc,
        "The project is organized into a layered architecture. Flutter owns the user experience, services "
        "own business logic and integration, Supabase owns identity and shared cloud data, SQLite provides "
        "offline resilience, and the Flask backend protects AI/RAG operations behind verified Supabase sessions.",
    )
    bullets(
        doc,
        [
            "Presentation layer: screens and widgets for dashboard, assets, manuals, AI assistant, collaboration, and settings.",
            "Service layer: authentication, database access, AI, RAG, synchronization, notifications, and collaboration.",
            "Persistence layer: local SQLite cache plus Supabase tables and storage.",
            "Backend layer: Flask endpoints that validate JWTs before executing AI or manual-retrieval workflows.",
        ],
    )

    para(doc, "2. Frontend Application Flow", "Heading 1")
    para(
        doc,
        "The app starts in main.dart, initializes Supabase, checks the cached session, and routes the user "
        "to either LoginScreen or HomeScreen. Once authenticated, HomeScreen becomes the hub for dashboard, "
        "asset analytics, manuals, AI assistant, collaboration, profile, and settings features.",
    )

    para(doc, "3. Authentication and Session Handling", "Heading 1")
    add_diagram(doc, "Figure 2. Google authentication flow", diagrams["auth"])
    para(
        doc,
        "Authentication is centralized in AuthService. Email/password and Google Sign-In both resolve into "
        "Supabase sessions, which means the rest of the app and backend can trust one session model. For Google "
        "auth, the app obtains an ID token from Google and exchanges it through Supabase using signInWithIdToken.",
    )
    bullets(
        doc,
        [
            "Supabase remains the source of truth for current user and session state.",
            "User profiles are synchronized into the users table after successful login or registration.",
            "Cached sessions are validated at startup so protected screens are not shown to invalid sessions.",
            "The backend accepts only Supabase JWT bearer tokens for protected AI/RAG operations.",
        ],
    )

    para(doc, "4. Android Google Auth Release Reliability", "Heading 1")
    add_diagram(doc, "Figure 3. Release signing and Google OAuth chain", diagrams["deployment"])
    para(
        doc,
        "Installed Android APKs require the package name, signing certificate fingerprint, Firebase Android "
        "OAuth client, google-services.json, and Supabase Google provider credentials to agree. This is the "
        "critical production step that prevents Google Sign-In from failing after installation on a real device.",
    )
    bullets(
        doc,
        [
            "Use a stable release keystore for production APKs.",
            "Register the release SHA-1 and SHA-256 fingerprints in Firebase/Google Cloud.",
            "Download the updated google-services.json after changing fingerprints.",
            "Use the matching Web OAuth client ID and secret in Supabase Auth provider settings.",
        ],
    )

    para(doc, "5. Data and Offline-First Strategy", "Heading 1")
    add_diagram(doc, "Figure 4. Data architecture", diagrams["data"])
    para(
        doc,
        "DatabaseHelper acts as the main gateway between UI workflows and persistence. The app can read and "
        "write from SQLite for local resilience while using Supabase tables and storage for authenticated cloud "
        "records, inventory images, manuals, and shared maintenance data.",
    )
    bullets(
        doc,
        [
            "SQLite gives the technician fast local access and offline fallback.",
            "Supabase tables provide cloud-backed assets, spare parts, service logs, manuals, and user data.",
            "Supabase Storage stores binary assets such as inventory images and manual PDFs.",
            "Fallback logic keeps clinical workflows usable even when the network is unstable.",
        ],
    )

    para(doc, "6. AI Assistant and RAG Workflow", "Heading 1")
    add_diagram(doc, "Figure 5. AI/RAG request lifecycle", diagrams["rag"])
    para(
        doc,
        "The AI assistant uses Gemini for diagnostic support and a backend RAG service for source-backed manual "
        "answers. The Flutter app sends the current Supabase access token to the Flask backend, where the token "
        "is verified before any retrieval, indexing, or AI generation work is allowed.",
    )
    bullets(
        doc,
        [
            "Manual questions are routed through RagApiService to the Flask backend.",
            "The backend verifies Supabase authentication before handling the request.",
            "Relevant manual chunks and source metadata are retrieved before response generation.",
            "Gemini generates the final answer using retrieved maintenance context.",
        ],
    )

    para(doc, "7. Crucial Modules", "Heading 1")
    component_table(doc)

    para(doc, "8. Security Model", "Heading 1")
    para(
        doc,
        "The security model is built around authenticated sessions and server-side verification. Client-side "
        "state is never enough for sensitive backend operations. Supabase row-level security and storage policies "
        "limit access to authenticated users and user-scoped paths, while Flask protects AI and RAG endpoints with "
        "JWT validation.",
    )
    bullets(
        doc,
        [
            "Supabase Auth controls application sessions.",
            "Backend APIs verify bearer tokens before executing protected operations.",
            "Storage uploads are scoped to authenticated user folders.",
            "Secrets and signing keys are kept out of source control through local config files.",
        ],
    )

    para(doc, "9. Why This Design Is Robust", "Heading 1")
    para(
        doc,
        "The project follows an industry-style split of responsibilities: Flutter is focused on interaction, "
        "services handle integration and domain logic, Supabase centralizes identity and data, SQLite protects "
        "field usability, and Flask keeps expensive AI workflows behind server-side authorization. This makes the "
        "system maintainable, secure, and resilient for real clinical maintenance environments.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "TitleCustom",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=29,
            textColor=colors.HexColor("#0B2A5B"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "SubtitleCustom",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#5B6770"),
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "H1Custom",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "H2Custom",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=9,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "BodyCustom",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#1F2B37"),
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "SmallCustom",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#1F2B37"),
        ),
    }


def pdf_bullets(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["body"]), leftIndent=12) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=7,
        bulletOffsetY=1,
    )


def add_pdf_diagram(story, title, image_path, styles):
    story.append(Paragraph(title, styles["h2"]))
    story.append(PdfImage(str(image_path), width=6.55 * inch, height=3.68 * inch))
    story.append(Spacer(1, 8))


def header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#5B6770"))
    canvas.drawString(0.72 * inch, 10.35 * inch, "Pulse Architecture Documentation")
    canvas.drawRightString(7.8 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    diagrams = build_diagrams()
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.7 * inch,
    )
    story = []

    story.append(Paragraph("Pulse Clinical Equipment Maintenance Platform", styles["title"]))
    story.append(Paragraph("Architecture and Crucial Implementation Documentation", styles["subtitle"]))
    story.append(
        Paragraph(
            "Pulse is a Flutter-based clinical equipment maintenance platform for biomedical technicians. "
            "It combines mobile-first workflows, authenticated cloud services, offline storage, AI diagnostics, "
            "manual retrieval, and collaboration tools into one operational system.",
            styles["body"],
        )
    )
    add_pdf_diagram(story, "Figure 1. System architecture at a glance", diagrams["high_level"], styles)

    story.append(Paragraph("1. Architecture Overview", styles["h1"]))
    story.append(
        Paragraph(
            "The project is organized into a layered architecture. Flutter owns the user experience, services "
            "own business logic and integration, Supabase owns identity and shared cloud data, SQLite provides "
            "offline resilience, and the Flask backend protects AI/RAG operations behind verified Supabase sessions.",
            styles["body"],
        )
    )
    story.append(
        pdf_bullets(
            [
                "Presentation layer: screens and widgets for dashboard, assets, manuals, AI assistant, collaboration, and settings.",
                "Service layer: authentication, database access, AI, RAG, synchronization, notifications, and collaboration.",
                "Persistence layer: local SQLite cache plus Supabase tables and storage.",
                "Backend layer: Flask endpoints that validate JWTs before executing AI or manual-retrieval workflows.",
            ],
            styles,
        )
    )

    story.append(Paragraph("2. Frontend Application Flow", styles["h1"]))
    story.append(
        Paragraph(
            "The app starts in main.dart, initializes Supabase, checks the cached session, and routes the user "
            "to either LoginScreen or HomeScreen. Once authenticated, HomeScreen becomes the hub for dashboard, "
            "asset analytics, manuals, AI assistant, collaboration, profile, and settings features.",
            styles["body"],
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("3. Authentication and Session Handling", styles["h1"]))
    add_pdf_diagram(story, "Figure 2. Google authentication flow", diagrams["auth"], styles)
    story.append(
        Paragraph(
            "Authentication is centralized in AuthService. Email/password and Google Sign-In both resolve into "
            "Supabase sessions, which means the rest of the app and backend can trust one session model. For Google "
            "auth, the app obtains an ID token from Google and exchanges it through Supabase using signInWithIdToken.",
            styles["body"],
        )
    )
    story.append(
        pdf_bullets(
            [
                "Supabase remains the source of truth for current user and session state.",
                "User profiles are synchronized into the users table after successful login or registration.",
                "Cached sessions are validated at startup so protected screens are not shown to invalid sessions.",
                "The backend accepts only Supabase JWT bearer tokens for protected AI/RAG operations.",
            ],
            styles,
        )
    )

    story.append(Paragraph("4. Android Google Auth Release Reliability", styles["h1"]))
    add_pdf_diagram(story, "Figure 3. Release signing and Google OAuth chain", diagrams["deployment"], styles)
    story.append(
        Paragraph(
            "Installed Android APKs require the package name, signing certificate fingerprint, Firebase Android "
            "OAuth client, google-services.json, and Supabase Google provider credentials to agree. This is the "
            "critical production step that prevents Google Sign-In from failing after installation on a real device.",
            styles["body"],
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("5. Data and Offline-First Strategy", styles["h1"]))
    add_pdf_diagram(story, "Figure 4. Data architecture", diagrams["data"], styles)
    story.append(
        Paragraph(
            "DatabaseHelper acts as the main gateway between UI workflows and persistence. The app can read and "
            "write from SQLite for local resilience while using Supabase tables and storage for authenticated cloud "
            "records, inventory images, manuals, and shared maintenance data.",
            styles["body"],
        )
    )
    story.append(
        pdf_bullets(
            [
                "SQLite gives the technician fast local access and offline fallback.",
                "Supabase tables provide cloud-backed assets, spare parts, service logs, manuals, and user data.",
                "Supabase Storage stores binary assets such as inventory images and manual PDFs.",
                "Fallback logic keeps clinical workflows usable even when the network is unstable.",
            ],
            styles,
        )
    )

    story.append(Paragraph("6. AI Assistant and RAG Workflow", styles["h1"]))
    add_pdf_diagram(story, "Figure 5. AI/RAG request lifecycle", diagrams["rag"], styles)
    story.append(
        Paragraph(
            "The AI assistant uses Gemini for diagnostic support and a backend RAG service for source-backed manual "
            "answers. The Flutter app sends the current Supabase access token to the Flask backend, where the token "
            "is verified before any retrieval, indexing, or AI generation work is allowed.",
            styles["body"],
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("7. Crucial Modules", styles["h1"]))
    data = [
        ["Layer", "Main Files", "Responsibility"],
        ["App shell", "main.dart", "Initialize Supabase, validate session, route to login or home."],
        ["Auth", "auth_service.dart", "Email auth, Google auth, password reset, session cleanup, user profile sync."],
        ["Data", "database_helper.dart", "SQLite cache, Supabase CRUD, storage uploads, offline fallback."],
        ["AI/RAG", "rag_api_service.dart, backend/app.py, backend/rag_service.py", "Protected manual retrieval and Gemini integration."],
        ["Collaboration", "chat_service.dart, google_chat_service.dart", "Team communication, meetings, and Google Chat notifications."],
        ["Operations", "notification_service.dart, sync_service.dart", "Local notifications, background behavior, sync strategy."],
    ]
    wrapped = [[Paragraph(str(cell), styles["small"]) for cell in row] for row in data]
    table = Table(wrapped, colWidths=[1.25 * inch, 2.0 * inch, 3.15 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F1FA")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2A5B")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D8DEE6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)

    story.append(Paragraph("8. Security Model", styles["h1"]))
    story.append(
        Paragraph(
            "The security model is built around authenticated sessions and server-side verification. Client-side "
            "state is never enough for sensitive backend operations. Supabase row-level security and storage policies "
            "limit access to authenticated users and user-scoped paths, while Flask protects AI and RAG endpoints with JWT validation.",
            styles["body"],
        )
    )
    story.append(
        pdf_bullets(
            [
                "Supabase Auth controls application sessions.",
                "Backend APIs verify bearer tokens before executing protected operations.",
                "Storage uploads are scoped to authenticated user folders.",
                "Secrets and signing keys are kept out of source control through local config files.",
            ],
            styles,
        )
    )

    story.append(Paragraph("9. Why This Design Is Robust", styles["h1"]))
    story.append(
        Paragraph(
            "The project follows an industry-style split of responsibilities: Flutter is focused on interaction, "
            "services handle integration and domain logic, Supabase centralizes identity and data, SQLite protects "
            "field usability, and Flask keeps expensive AI workflows behind server-side authorization. This makes the "
            "system maintainable, secure, and resilient for real clinical maintenance environments.",
            styles["body"],
        )
    )

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
    return PDF_PATH


if __name__ == "__main__":
    docx_path = build_docx()
    pdf_path = build_pdf()
    print(docx_path)
    print(pdf_path)
